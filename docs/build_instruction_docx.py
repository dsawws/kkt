# -*- coding: utf-8 -*-
"""Генерация Word-инструкции по деплою и обновлению сайта ККТ."""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "доки")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_PATH = os.path.join(
    OUT_DIR, "Инструкция_деплой_и_обновление_сайта_ККТ_v1.1.docx"
)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    sizes = {1: 16, 2: 14, 3: 13}
    for run in p.runs:
        set_run_font(run, size=sizes.get(level, 12), bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, size=12, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.75)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Важно. ")
    set_run_font(run, bold=True, size=11)
    run2 = p.add_run(text)
    set_run_font(run2, size=11, italic=True)
    return p


def page_break(doc):
    doc.add_page_break()


def centered(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ---- Титул ----
    for _ in range(3):
        doc.add_paragraph()
    centered(doc, "КРАСНОДАРСКИЙ КООПЕРАТИВНЫЙ ТЕХНИКУМ", size=14, bold=True)
    centered(doc, "Сайт new.kktbel.ru", size=12)
    for _ in range(2):
        doc.add_paragraph()
    centered(doc, "ИНСТРУКЦИЯ", size=22, bold=True)
    centered(doc, "по развёртыванию, обновлению фронтенда", size=14)
    centered(doc, "и безопасной работе с контентом и файлами", size=14)
    for _ in range(4):
        doc.add_paragraph()
    centered(doc, "Версия документа: 1.0", size=11)
    centered(doc, f"Дата: {date.today().strftime('%d.%m.%Y')}", size=11)
    centered(doc, "Платформа: Django + Gunicorn + Nginx + HTTPS", size=11)
    page_break(doc)

    # ---- Содержание ----
    add_heading_custom(doc, "СОДЕРЖАНИЕ", 1)
    toc_items = [
        "1. Назначение документа и область применения",
        "2. Общая архитектура сайта",
        "3. Что относится к фронтенду, а что — к данным и файлам",
        "4. Структура каталогов на сервере",
        "5. Переменные окружения и режим HTTPS",
        "6. Первичная установка (один раз)",
        "7. Штатное обновление сайта (рекомендуемый способ)",
        "8. Обновление только фронтенда без риска для документов",
        "9. Сборка статики (collectstatic) и роль Nginx",
        "10. Работа с контентом через панель администратора",
        "    10.1. Раздел «Сведения об организации»",
        "    10.2. Новости",
        "    10.3. Документы и PDF-файлы",
        "    10.4. Страницы, меню, таблицы",
        "11. CI/CD (GitHub Actions)",
        "12. Резервное копирование",
        "13. Типовые неисправности и диагностика",
        "14. Правила безопасности и запрещённые действия",
        "15. Чек-листы",
        "16. Приложения (команды и пути)",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0)
        set_run_font(p.add_run(item), size=12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.first_line_indent = Cm(0)
    set_run_font(
        p.add_run(
            "Нумерация разделов соответствует заголовкам в тексте документа. "
            "Документ рассчитан на печать и чтение постранично: каждый крупный "
            "раздел начинается с новой страницы."
        ),
        size=10,
        italic=True,
    )
    page_break(doc)

    # ---- 1 ----
    add_heading_custom(doc, "1. Назначение документа и область применения", 1)
    add_para(
        doc,
        "Настоящая инструкция описывает полный цикл сопровождения официального "
        "сайта Краснодарского кооперативного техникума (рабочий адрес "
        "продакшен-среды — https://new.kktbel.ru). Документ предназначен для "
        "администраторов сервера, разработчиков фронтенда и контент-менеджеров, "
        "работающих через панель управления /panel/.",
    )
    add_para(
        doc,
        "Цель инструкции — обеспечить безопасное обновление внешнего вида и "
        "программного кода сайта без повреждения загруженных документов, "
        "медиафайлов, базы данных и настроек бэкенда. Отдельно разъясняется, "
        "какие операции затрагивают только статику (CSS, JavaScript, изображения "
        "интерфейса), а какие — пользовательский контент организации.",
    )
    add_para(
        doc,
        "Инструкция опирается на принятую в проекте схему: исходный код хранится "
        "в системе контроля версий Git; на сервере выполняется git clone; статика "
        "собирается командой collectstatic; документы и изображения, загруженные "
        "через панель, лежат в каталоге media и не входят в репозиторий. При "
        "соблюдении этой схемы обновление «фронта» и наполнение сайта контентом "
        "не мешают друг другу.",
    )
    add_para(
        doc,
        "Рекомендуется хранить настоящий файл вместе с проектной документацией "
        "(каталог «доки») и актуализировать его при существенных изменениях "
        "процедуры деплоя или структуры каталогов на сервере.",
    )
    page_break(doc)

    # ---- 2 ----
    add_heading_custom(doc, "2. Общая архитектура сайта", 1)
    add_para(
        doc,
        "Сайт построен на фреймворке Django. В продакшене запросы пользователей "
        "принимает веб-сервер Nginx. Статические файлы и загруженные документы "
        "Nginx отдаёт напрямую с диска. Динамические страницы (HTML, формируемый "
        "Django), панель администратора и API проксируются на приложение "
        "Gunicorn, которое слушает 127.0.0.1:8041; Nginx на HTTP :8042 проксирует на бэкенд и не "
        "доступно извне напрямую.",
    )
    add_para(doc, "Схема обработки запроса в упрощённом виде:")
    add_bullet(doc, "Браузер → HTTPS (порт 443) → Nginx.")
    add_bullet(
        doc,
        "Запрос к /static/… → файлы из django_admin/staticfiles/ "
        "(результат collectstatic).",
    )
    add_bullet(
        doc,
        "Запрос к /media/… → файлы из django_admin/media/ "
        "(загрузки панели, PDF, изображения новостей и т. п.).",
    )
    add_bullet(
        doc,
        "Запрос к /uploads/… → legacy-каталог старого сайта (при наличии).",
    )
    add_bullet(
        doc,
        "Остальные URL → Gunicorn (Django): страницы CMS, новости /news/, "
        "панель /panel/, загрузка через CKEditor.",
    )
    add_para(
        doc,
        "Управление процессом Gunicorn выполняется через systemd-сервис kktbel. "
        "Секреты и параметры продакшена задаются в файле /var/www/html/deploy/kktbel.env и "
        "подхватываются сервисом. Шаблон unit-файла находится в "
        "deploy/kktbel.service, шаблон Nginx — в deploy/nginx-new.kktbel.ru.conf.",
    )
    add_note(
        doc,
        "Сайт работает по HTTPS. Nginx должен передавать заголовок "
        "X-Forwarded-Proto, а в окружении Django включены secure-cookies, "
        "редирект на HTTPS и HSTS.",
    )
    page_break(doc)

    # ---- 3 ----
    add_heading_custom(
        doc, "3. Что относится к фронтенду, а что — к данным и файлам", 1
    )
    add_para(
        doc,
        "Ключевой принцип безопасного обновления: не смешивать «код интерфейса» "
        "и «данные организации». Ниже приведено разделение, которого следует "
        "придерживаться всегда — и при ручном деплое, и при автоматическом "
        "через GitHub Actions.",
    )

    add_heading_custom(
        doc, "3.1. Фронтенд и код приложения (обновляются из Git)", 2
    )
    add_bullet(
        doc,
        "Исходники стилей и скриптов: django_admin/static/ "
        "(style.css, script.js, panel/* и др.).",
    )
    add_bullet(doc, "HTML-шаблоны: django_admin/templates/.")
    add_bullet(doc, "Логика Django: django_admin/cms/, techcollege_admin/.")
    add_bullet(doc, "Скрипты деплоя и примеры конфигов: deploy/.")
    add_bullet(
        doc,
        "После выкладки на сервере заново собирается каталог staticfiles/ — "
        "именно его видит Nginx по адресу /static/.",
    )

    add_heading_custom(
        doc,
        "3.2. Данные и файлы (не из Git; collectstatic их не трогает)",
        2,
    )
    add_bullet(
        doc,
        "База данных SQLite: django_admin/db.sqlite3 — страницы, новости, меню, "
        "записи о документах, настройки главной.",
    )
    add_bullet(
        doc,
        "Медиафайлы: django_admin/media/ — PDF, загрузки CKEditor, изображения "
        "новостей, баннеры.",
    )
    add_bullet(doc, "Legacy uploads: uploads/ — старые абсолютные ссылки.")
    add_bullet(
        doc,
        "Секреты: /var/www/html/deploy/kktbel.env — ключи и хосты (вне репозитория).",
    )
    add_para(
        doc,
        "В .gitignore проекта явно исключены db.sqlite3, media/ и staticfiles/. "
        "Поэтому обычный git pull и скрипт remote-update.sh не удаляют документы "
        "и не подменяют базу данных локальной копией разработчика.",
    )
    add_note(
        doc,
        "Никогда не копируйте на сервер zip-архив «всего проекта с компьютера», "
        "если в нём есть пустые или тестовые media/ и db.sqlite3 — так можно "
        "случайно затереть боевые данные.",
    )
    page_break(doc)

    # ---- 4 ----
    add_heading_custom(doc, "4. Структура каталогов на сервере", 1)
    add_para(
        doc,
        "Типовое расположение (если при установке задан другой путь, "
        "скорректируйте команды; в скриптах по умолчанию используется "
        "/var/www/html):",
    )
    add_code(
        doc,
        "/var/www/html/                    # корень git-репозитория\n"
        "├── deploy/                       # скрипты и конфиги деплоя\n"
        "│   ├── collectstatic.sh\n"
        "│   ├── remote-update.sh\n"
        "│   ├── kktbel.service\n"
        "│   ├── kktbel.env.example\n"
        "│   └── nginx-new.kktbel.ru.conf\n"
        "├── django_admin/                 # приложение Django\n"
        "│   ├── static/                   # ИСХОДНИКИ фронта (в Git)\n"
        "│   ├── staticfiles/              # СБОРКА для Nginx (не в Git)\n"
        "│   ├── media/                    # ДОКУМЕНТЫ И ЗАГРУЗКИ (не в Git)\n"
        "│   ├── db.sqlite3                # БАЗА (не в Git)\n"
        "│   ├── templates/\n"
        "│   ├── cms/\n"
        "│   └── venv/\n"
        "└── uploads/                      # legacy (по необходимости)",
    )
    add_para(doc, "Дополнительно вне дерева сайта:")
    add_bullet(doc, "/var/www/html/deploy/kktbel.env — переменные окружения продакшена.")
    add_bullet(doc, "/etc/systemd/system/kktbel.service — unit Gunicorn.")
    add_bullet(
        doc,
        "/etc/nginx/sites-available/new.kktbel.ru — виртуальный хост Nginx.",
    )
    page_break(doc)

    # ---- 5 ----
    add_heading_custom(doc, "5. Переменные окружения и режим HTTPS", 1)
    add_para(
        doc,
        "Шаблон переменных: deploy/kktbel.env.example. Рабочий файл на сервере: "
        "/var/www/html/deploy/kktbel.env (права 640, readable для www-data (gunicorn)). "
        "Минимально необходимый набор для HTTPS-продакшена:",
    )
    add_code(
        doc,
        "DJANGO_DEBUG=0\n"
        "DJANGO_SECRET_KEY=<длинная случайная строка>\n"
        "DJANGO_ALLOWED_HOSTS=new.kktbel.ru,kktbel.ru,www.kktbel.ru\n"
        "DJANGO_CSRF_TRUSTED_ORIGINS=https://new.kktbel.ru,"
        "https://kktbel.ru,https://www.kktbel.ru\n"
        "DJANGO_SECURE_SSL_REDIRECT=1\n"
        "DJANGO_SESSION_COOKIE_SECURE=1\n"
        "DJANGO_CSRF_COOKIE_SECURE=1\n"
        "DJANGO_SECURE_HSTS_SECONDS=31536000",
    )
    add_para(doc, "Пояснения к параметрам:")
    add_bullet(
        doc,
        "DJANGO_DEBUG=0 — отключает отладочные страницы ошибок и включает "
        "продакшен-режим безопасности.",
    )
    add_bullet(
        doc,
        "DJANGO_SECRET_KEY — обязателен; значение, начинающееся с "
        "django-insecure-, в продакшене запрещено приложением.",
    )
    add_bullet(
        doc,
        "CSRF_TRUSTED_ORIGINS — указывайте только https://… для рабочих доменов.",
    )
    add_bullet(
        doc,
        "SECURE_SSL_REDIRECT и secure-cookies — корректная работа сессий и форм "
        "за HTTPS.",
    )
    add_bullet(
        doc,
        "HSTS — браузер запоминает обязательность HTTPS; включать, когда "
        "сертификат стабилен. Опцию для всех поддоменов включайте только если "
        "все они тоже на SSL.",
    )
    add_para(
        doc,
        "Nginx обязан проксировать X-Forwarded-Proto $scheme, иначе Django не "
        "«увидит» HTTPS за reverse proxy. В шаблоне nginx настроено: HTTP "
        "(порт 80) перенаправляет на HTTPS; приложение и статика обслуживаются "
        "на 443. Пути к сертификатам в шаблоне — стандартные для Let’s Encrypt; "
        "при ином CA скорректируйте ssl_certificate и ssl_certificate_key.",
    )
    page_break(doc)

    # ---- 6 ----
    add_heading_custom(doc, "6. Первичная установка (один раз)", 1)
    add_para(
        doc,
        "Ниже — порядок первой установки на чистый сервер. Если сайт уже "
        "развёрнут, используйте раздел как справочник и сверяйте отличия "
        "конфигурации.",
    )
    add_num(
        doc,
        "Клонировать репозиторий в /var/www/html именно через git clone "
        "(не распаковкой zip без каталога .git).",
    )
    add_num(doc, "Создать виртуальное окружение и установить зависимости:")
    add_code(
        doc,
        "cd /var/www/html/django_admin\n"
        "python3 -m venv venv && source venv/bin/activate\n"
        "pip install -r requirements.txt",
    )
    add_num(doc, "Скопировать и заполнить секреты:")
    add_code(
        doc,
        "sudo cp /var/www/html/deploy/kktbel.env.example /var/www/html/deploy/kktbel.env\n"
        "sudo chmod 600 /var/www/html/deploy/kktbel.env\n"
        "sudo nano /var/www/html/deploy/kktbel.env",
    )
    add_num(
        doc,
        "Выполнить миграции, собрать статику, создать суперпользователя:",
    )
    add_code(
        doc,
        "set -a && source /var/www/html/deploy/kktbel.env && set +a\n"
        "export DJANGO_DEBUG=0\n"
        "python manage.py migrate\n"
        "bash /var/www/html/deploy/collectstatic.sh\n"
        "python manage.py createsuperuser",
    )
    add_num(
        doc,
        "Установить и запустить systemd-сервис kktbel "
        "(скопировать unit, daemon-reload, enable --now).",
    )
    add_num(
        doc,
        "Подключить конфиг Nginx, проверить nginx -t, выполнить reload. "
        "Убедиться, что пути к SSL-сертификатам соответствуют фактическим.",
    )
    add_num(
        doc,
        "Настроить права: код читается процессом gunicorn (обычно www-data); "
        "каталог media/ доступен на запись пользователю процесса Django и на "
        "чтение Nginx.",
    )
    add_note(
        doc,
        "После установки сохраните отдельно резервную копию db.sqlite3 и "
        "каталога media/ на носитель вне сервера сайта.",
    )
    page_break(doc)

    # ---- 7 ----
    add_heading_custom(
        doc, "7. Штатное обновление сайта (рекомендуемый способ)", 1
    )
    add_para(
        doc,
        "Рекомендуемый и наиболее безопасный способ обновления продакшена — "
        "скрипт deploy/remote-update.sh. Он обновляет код из Git, ставит "
        "зависимости, делает резервную копию SQLite, применяет миграции, "
        "пересобирает статику и перезапускает Gunicorn. Каталог media/ скрипт "
        "не очищает и не заменяет.",
    )
    add_para(doc, "Запуск на сервере из корня репозитория:")
    add_code(doc, "bash /var/www/html/deploy/remote-update.sh main")
    add_para(
        doc,
        "Вместо main можно указать другую ветку или тег. Последовательность "
        "действий скрипта:",
    )
    add_num(
        doc,
        "git fetch / checkout / reset --hard origin/<ref> — код приводится к "
        "состоянию удалённой ветки.",
    )
    add_num(doc, "pip install -r requirements.txt — обновление Python-зависимостей.")
    add_num(
        doc,
        "Копия db.sqlite3 → db.sqlite3.bak.<дата-время> — страховка перед "
        "миграциями.",
    )
    add_num(doc, "python manage.py migrate — схема БД без удаления контента.")
    add_num(
        doc,
        "bash deploy/collectstatic.sh — полная пересборка staticfiles для Nginx.",
    )
    add_num(doc, "systemctl restart kktbel — новый код в процессе Gunicorn.")
    add_num(
        doc,
        "Проверка доступности главной и ключевых static-файлов через curl.",
    )
    add_para(
        doc,
        "После обновления имеет смысл вручную открыть в браузере: главную "
        "страницу, раздел «Сведения об организации», /news/, одну страницу с "
        "PDF и панель /panel/ (вход).",
    )
    page_break(doc)

    # ---- 8 ----
    add_heading_custom(
        doc,
        "8. Обновление только фронтенда без риска для документов",
        1,
    )
    add_para(
        doc,
        "Если менялись только стили, скрипты или шаблоны (внешний вид и "
        "поведение интерфейса), достаточно доставить код и пересобрать статику. "
        "Документы в media/ и записи в базе при этом не затрагиваются — при "
        "условии, что вы не копируете эти каталоги вручную и не используете "
        "деструктивные команды очистки репозитория.",
    )

    add_heading_custom(
        doc, "8.1. Через полный штатный деплой (предпочтительно)", 2
    )
    add_para(
        doc,
        "Даже для «чисто фронтовых» правок удобнее использовать "
        "remote-update.sh: процедура стандартизирована, меньше шансов забыть "
        "collectstatic или restart. Миграции при отсутствии изменений схемы "
        "выполняются быстро и безопасно для данных.",
    )

    add_heading_custom(doc, "8.2. Узкий сценарий «только статика»", 2)
    add_para(
        doc,
        "Если код уже обновлён (git pull выполнен), а нужно лишь пересобрать "
        "фронт:",
    )
    add_code(
        doc,
        "cd /var/www/html/django_admin\n"
        "source venv/bin/activate\n"
        "set -a && source /var/www/html/deploy/kktbel.env && set +a\n"
        "bash /var/www/html/deploy/collectstatic.sh\n"
        "sudo systemctl restart kktbel",
    )
    add_para(
        doc,
        "Скрипт collectstatic.sh вызывает collectstatic --noinput --clear и "
        "затем check_static. Флаг --clear очищает только staticfiles/, не media/ "
        "и не базу данных.",
    )

    add_heading_custom(doc, "8.3. Локальная разработка", 2)
    add_bullet(doc, "Править файлы в django_admin/static/ и templates/.")
    add_bullet(
        doc,
        "Проверять через runserver (DEBUG=1); для локальной проверки "
        "collectstatic обычно не обязателен.",
    )
    add_bullet(
        doc,
        "Коммитить и пушить в Git; на сервер не копировать media/ и db.sqlite3 "
        "с рабочей машины.",
    )
    add_bullet(
        doc,
        "После выкладки на прод обновить страницу с принудительным сбросом кэша "
        "(Ctrl+F5).",
    )

    add_heading_custom(
        doc, "8.4. Что гарантированно нельзя делать при обновлении фронта", 2
    )
    add_bullet(doc, "Удалять или «синхронизировать поверх» каталог media/.")
    add_bullet(doc, "Заливать локальный db.sqlite3 на сервер.")
    add_bullet(
        doc,
        "Править файлы напрямую в staticfiles/ — при следующем collectstatic "
        "правки пропадут; исходники правятся в static/.",
    )
    add_bullet(
        doc,
        "Отключать location /media/ в Nginx — список документов из БД останется, "
        "но PDF начнут отдавать 404.",
    )
    page_break(doc)

    # ---- 9 ----
    add_heading_custom(
        doc, "9. Сборка статики (collectstatic) и роль Nginx", 1
    )
    add_para(
        doc,
        "Django хранит исходники статики в нескольких местах: каталог проекта "
        "static/, статика приложений, пакеты (CKEditor, MPTT, admin). Команда "
        "collectstatic собирает всё в единый STATIC_ROOT = "
        "django_admin/staticfiles/. Nginx обслуживает только этот каталог по "
        "URL /static/.",
    )
    add_para(
        doc,
        "Не добавляйте в Nginx отдельные location для ckeditor, mptt, admin или "
        "panel: после корректного collectstatic эти файлы уже лежат внутри "
        "staticfiles/.",
    )
    add_para(doc, "Проверка после сборки:")
    add_code(
        doc,
        "python manage.py check_static\n"
        "curl -I https://new.kktbel.ru/static/style.css\n"
        "curl -I https://new.kktbel.ru/static/panel/admin.css\n"
        "curl -I https://new.kktbel.ru/static/ckeditor/ckeditor/ckeditor.js",
    )
    add_para(
        doc,
        "Если check_static сообщает об отсутствии файлов — проблема в "
        "зависимостях pip или в настройках STATICFILES, а не в media и не в "
        "документах организации.",
    )
    page_break(doc)

    # ---- 10 ----
    add_heading_custom(
        doc, "10. Работа с контентом через панель администратора", 1
    )
    add_para(
        doc,
        "Контент сайта (тексты, таблицы, новости, PDF) вносится через панель "
        "https://new.kktbel.ru/panel/ и сохраняется в базе и в media/. Это не "
        "требует деплоя кода и не затрагивает фронтенд-репозиторий. Деплой "
        "нужен только когда меняется программный код или шаблоны.",
    )

    add_heading_custom(doc, "10.1. Раздел «Сведения об организации»", 2)
    add_para(
        doc,
        "В боковом меню панели есть пункт «Сведения об организации» "
        "(/panel/organization/). Там перечислены подстраницы раздела: статус "
        "текста, число документов, кнопки перехода в редактор, добавление файла "
        "и просмотр на сайте. Недостающие служебные подстраницы создаются "
        "автоматически; уже заполненный HTML при этом не затирается.",
    )
    add_para(
        doc,
        "Рекомендуемый порядок заполнения: открыть нужную подстраницу → ввести "
        "текст и таблицы в редакторе → при необходимости вставить "
        "переиспользуемую таблицу кнопкой «Таблица» → прикрепить PDF через "
        "«+ Документ» с привязкой к этой странице. Документы, привязанные к "
        "странице, отображаются внизу публичной страницы автоматически.",
    )

    add_heading_custom(doc, "10.2. Новости", 2)
    add_para(
        doc,
        "Лента новостей обслуживается отдельными маршрутами /news/ и "
        "/news/<slug>/, а не произвольной CMS-страницей. В меню пункт "
        "«Новости» должен вести на /news/. Устаревшие адреса вида "
        "/page/novosti-test/ и /page/news/ перенаправляются на ленту. Создание "
        "и редактирование — раздел «Новости» в панели. На главной выводятся "
        "последние опубликованные материалы.",
    )

    add_heading_custom(doc, "10.3. Документы и PDF-файлы", 2)
    add_para(
        doc,
        "Файл документа сохраняется на диск в media/documents/…; в базе "
        "хранится запись Document с привязкой к странице (поле page). Список на "
        "сайте формируется Django из БД; скачивание и открытие идут через "
        "Nginx по пути /media/.",
    )
    add_para(
        doc,
        "Если на странице нет списка — документ привязан к другой странице или "
        "неактивен. Если список есть, а «Открыть» даёт 404 — отсутствует файл "
        "на диске или не настроен location /media/. Команда collectstatic к "
        "документам отношения не имеет.",
    )

    add_heading_custom(doc, "10.4. Страницы, меню, таблицы", 2)
    add_para(
        doc,
        "Обычные страницы редактируются в разделе «Страницы». Флаг "
        "«Показывать в меню» синхронизирует пункт MenuItem. Переиспользуемые "
        "HTML-таблицы создаются в разделе «Таблицы» и вставляются в текст "
        "страницы. Меню на сайте строится из активных корневых пунктов с "
        "дедупликацией одинаковых названий (предпочтение пункту со связанной "
        "страницей).",
    )
    page_break(doc)

    # ---- 11 ----
    add_heading_custom(doc, "11. CI/CD (GitHub Actions)", 1)
    add_para(doc, "В репозитории настроены два workflow.")
    add_bullet(
        doc,
        "CI — автоматически на push и pull request: manage.py check, тесты, "
        "collectstatic, smoke- и security-проверки. Красный пайплайн — сигнал "
        "не выкатывать изменения в прод.",
    )
    add_bullet(
        doc,
        "Deploy — только вручную (Actions → Deploy → Run workflow): SSH на "
        "сервер и выполнение remote-update.sh.",
    )
    add_para(
        doc,
        "Необходимые Secrets репозитория: DEPLOY_HOST, DEPLOY_USER, "
        "DEPLOY_SSH_KEY; опционально DEPLOY_PORT и DEPLOY_PATH. Подробная "
        "пошаговая настройка описана в файле deploy/CI-CD.md. Environment "
        "production можно защитить обязательным подтверждением ревьюера перед "
        "деплоем.",
    )
    add_para(
        doc,
        "Условие корректного деплоя: каталог на сервере — git clone того же "
        "remote, что и GitHub, а не «ручная копия файлов» без истории Git.",
    )
    page_break(doc)

    # ---- 12 ----
    add_heading_custom(doc, "12. Резервное копирование", 1)
    add_para(
        doc,
        "Перед крупными изменениями и регулярно (например, ежедневно или "
        "еженедельно) сохраняйте:",
    )
    add_bullet(
        doc, "django_admin/db.sqlite3 — вся структура и тексты сайта."
    )
    add_bullet(
        doc, "django_admin/media/ — все загруженные через панель файлы."
    )
    add_bullet(
        doc,
        "При необходимости — uploads/ и актуальный /var/www/html/deploy/kktbel.env "
        "(отдельно, с ограничением доступа).",
    )
    add_para(
        doc,
        "Скрипт remote-update.sh перед migrate уже создаёт db.sqlite3.bak.* в "
        "том же каталоге — это краткосрочная страховка, не замена полноценного "
        "бэкапа на другой носитель или в облачное хранилище.",
    )
    add_code(
        doc,
        "# Пример ручного бэкапа на сервере\n"
        "TS=$(date +%Y%m%d-%H%M%S)\n"
        "cp -a /var/www/html/django_admin/db.sqlite3 "
        "/backup/kkt-db-$TS.sqlite3\n"
        "tar -czf /backup/kkt-media-$TS.tar.gz "
        "-C /var/www/html/django_admin media",
    )
    page_break(doc)

    # ---- 13 ----
    add_heading_custom(doc, "13. Типовые неисправности и диагностика", 1)

    add_heading_custom(
        doc, "13.1. Стили «старые» или пропали после выкладки", 2
    )
    add_bullet(
        doc,
        "Не выполнен collectstatic или не перезапущен gunicorn "
        "(для HTML из шаблонов важен restart).",
    )
    add_bullet(
        doc, "Кэш браузера — проверить в режиме инкогнито или Ctrl+F5."
    )
    add_bullet(
        doc,
        "Nginx смотрит не на staticfiles/ — сверить alias в конфиге.",
    )
    add_code(
        doc,
        "bash /var/www/html/deploy/collectstatic.sh\n"
        "sudo systemctl restart kktbel\n"
        "curl -I https://new.kktbel.ru/static/style.css",
    )

    add_heading_custom(doc, "13.2. Документ виден в списке, PDF — 404", 2)
    add_bullet(doc, "Проверить наличие файла в media/documents/…")
    add_bullet(
        doc, "Проверить location /media/ и права чтения для Nginx."
    )
    add_bullet(doc, "Не путать с /static/ и командой collectstatic.")

    add_heading_custom(
        doc, "13.3. Документов нет в списке на странице", 2
    )
    add_bullet(
        doc, "В панели у документа указана другая «Страница»."
    )
    add_bullet(doc, "Снят флаг «Активен».")

    add_heading_custom(
        doc, "13.4. Новости открываются как 404 /page/…", 2
    )
    add_bullet(
        doc,
        "Пункт меню должен вести на /news/, не на несуществующую CMS-страницу.",
    )
    add_bullet(
        doc,
        "Проверить /news/ напрямую; при необходимости исправить меню в панели.",
    )

    add_heading_custom(
        doc, "13.5. Ошибки CSRF или зацикленные редиректы", 2
    )
    add_bullet(
        doc,
        "В CSRF_TRUSTED_ORIGINS должны быть https-origins рабочих доменов.",
    )
    add_bullet(doc, "Nginx передаёт X-Forwarded-Proto корректно.")
    add_bullet(
        doc,
        "При работе за reverse proxy не отключайте SECURE_PROXY_SSL_HEADER "
        "в продакшене.",
    )

    add_heading_custom(doc, "13.6. Сервис не поднимается", 2)
    add_code(
        doc,
        "sudo systemctl status kktbel\n"
        "journalctl -u kktbel -n 100 --no-pager\n"
        "# частая причина: пустой или неверный DJANGO_SECRET_KEY при DEBUG=0",
    )
    page_break(doc)

    # ---- 14 ----
    add_heading_custom(
        doc, "14. Правила безопасности и запрещённые действия", 1
    )
    add_bullet(
        doc, "Не хранить SECRET_KEY в Git и не публиковать /var/www/html/deploy/kktbel.env."
    )
    add_bullet(doc, "Не оставлять DEBUG=1 на продакшене.")
    add_bullet(
        doc,
        "Не использовать force push в main без крайней необходимости и "
        "согласования.",
    )
    add_bullet(
        doc,
        "Не выполнять на сервере git clean -fdx по корню проекта — можно "
        "удалить media и локальные артефакты.",
    )
    add_bullet(doc, "Не подменять боевую базу данных тестовой.")
    add_bullet(
        doc,
        "Ограничить SSH-доступ; для деплоя — отдельный пользователь с "
        "минимально нужным sudo на restart kktbel.",
    )
    add_bullet(
        doc,
        "После настройки HTTPS не отключать secure-cookies без веской причины.",
    )
    add_para(
        doc,
        "Рекомендуемые права: владелец дерева кода — пользователь деплоя; "
        "группа www-data; media/ доступен на запись www-data; staticfiles/ "
        "доступен на чтение Nginx.",
    )
    page_break(doc)

    # ---- 15 ----
    add_heading_custom(doc, "15. Чек-листы", 1)

    add_heading_custom(doc, "15.1. Перед выкладкой фронта", 2)
    add_bullet(doc, "Изменения закоммичены и попали в нужную ветку Git.")
    add_bullet(doc, "Локально проверены ключевые страницы.")
    add_bullet(
        doc,
        "Есть свежий бэкап db.sqlite3 и media/ (особенно перед миграциями).",
    )
    add_bullet(doc, "CI зелёный (если используется GitHub Actions).")

    add_heading_custom(doc, "15.2. После выкладки", 2)
    add_bullet(doc, "https://new.kktbel.ru/ открывается.")
    add_bullet(doc, "/static/style.css отдаёт код ответа 200.")
    add_bullet(doc, "Раздел сведений и новости открываются.")
    add_bullet(doc, "Открывается тестовый PDF из media.")
    add_bullet(doc, "Вход в /panel/ работает.")
    add_bullet(doc, "systemctl status kktbel — active (running).")

    add_heading_custom(doc, "15.3. Только контент (без деплоя)", 2)
    add_bullet(doc, "Права доступа в панели подтверждены.")
    add_bullet(doc, "Страница опубликована.")
    add_bullet(
        doc,
        "Документ привязан к верной странице, файл загружен, флаг «Активен» "
        "включён.",
    )
    add_bullet(
        doc, "Просмотр выполнен с публичного URL, не только из панели."
    )
    page_break(doc)

    # ---- 16 ----
    add_heading_custom(doc, "16. Приложения (команды и пути)", 1)

    add_heading_custom(doc, "16.1. Быстрые команды", 2)
    add_code(
        doc,
        "# Полное обновление\n"
        "bash /var/www/html/deploy/remote-update.sh main\n"
        "\n"
        "# Только статика\n"
        "bash /var/www/html/deploy/collectstatic.sh\n"
        "sudo systemctl restart kktbel\n"
        "\n"
        "# Статус\n"
        "sudo systemctl status kktbel\n"
        "sudo nginx -t && sudo systemctl reload nginx",
    )

    add_heading_custom(doc, "16.2. Важные URL", 2)
    add_bullet(doc, "Сайт: https://new.kktbel.ru/")
    add_bullet(doc, "Новости: https://new.kktbel.ru/news/")
    add_bullet(doc, "Панель: https://new.kktbel.ru/panel/")
    add_bullet(
        doc, "Хаб сведений: https://new.kktbel.ru/panel/organization/"
    )

    add_heading_custom(doc, "16.3. Краткая памятка «что трогаем»", 2)
    add_para(
        doc,
        "Обновляем фронт — Git, затем collectstatic и restart сервиса. "
        "Контент и PDF — только панель (база данных и media). Секреты — "
        "/var/www/html/deploy/kktbel.env. Не смешивайте эти три контура: тогда бэкенд и "
        "документы организации останутся целыми при любых обновлениях "
        "интерфейса.",
    )

    doc.add_paragraph()
    centered(doc, "— Конец документа —", size=11, italic=True)

    doc.save(OUT_PATH)
    print("SAVED", OUT_PATH)
    print("SIZE", os.path.getsize(OUT_PATH))


if __name__ == "__main__":
    build()
